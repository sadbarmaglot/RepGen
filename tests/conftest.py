import sys
from unittest.mock import MagicMock, patch

# Мокаем GCS storage client ДО импорта любых модулей проекта,
# чтобы не требовались реальные credentials при тестах.
mock_storage_client = MagicMock()
mock_bucket = MagicMock()
mock_storage_client.bucket.return_value = mock_bucket

# Патчим google.cloud.storage.Client чтобы gc_utils.py не падал при импорте
with patch.dict("os.environ", {"JWT_SECRET_KEY": "test-secret-key"}):
    patch("google.cloud.storage.Client", return_value=mock_storage_client).start()

import pytest


@pytest.fixture
def sample_docx_bytes():
    """Создаёт минимальный .docx в памяти."""
    from docx import Document
    import io

    doc = Document()
    doc.add_paragraph("Заголовок отчёта")
    doc.add_paragraph("Объект обследования: ул. Ленина, д. 1")
    doc.add_paragraph("")  # пустой параграф — должен пропуститься

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Параметр"
    table.cell(0, 1).text = "Значение"
    table.cell(1, 0).text = "Этажность"
    table.cell(1, 1).text = "3"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def sample_pdf_bytes():
    """Создаёт минимальный PDF в памяти через pymupdf."""
    import pymupdf

    doc = pymupdf.open()
    page = doc.new_page()
    # pymupdf insert_text с дефолтным шрифтом не рендерит кириллицу,
    # используем latin для теста парсинга.
    page.insert_text((72, 72), "Technical Report\nAddress: Moscow", fontsize=12)
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


@pytest.fixture
def empty_docx_bytes():
    """Пустой .docx без текста."""
    from docx import Document
    import io

    doc = Document()
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def reference_docs_dir(tmp_path):
    """Временная директория с reference-документами."""
    ref_dir = tmp_path / "reference_docs"
    ref_dir.mkdir()

    (ref_dir / "GOST_31937.txt").write_text(
        "ГОСТ 31937-2024 Обследование зданий", encoding="utf-8"
    )
    (ref_dir / "SP_20.txt").write_text(
        "СП 20.13330.2016 Нагрузки и воздействия", encoding="utf-8"
    )

    return ref_dir
