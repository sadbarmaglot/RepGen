import asyncio
import io
import logging
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument
import pymupdf

from common.gc_utils import documents_storage

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = {".docx", ".pdf"}
MAX_DOCUMENT_SIZE = 100 * 1024 * 1024  # 100 MB

DEFAULT_MODEL = "gpt-5.4"
REVIEW_MAX_TOKENS = 32768
REVIEW_TEMPERATURE = 0.2

AVAILABLE_MODELS = {
    "gpt-5.4": {"label": "GPT-5.4"},
    "gpt-5.4-mini": {"label": "GPT-5.4 Mini"},
}

REFERENCE_DOCS_DIR = Path(__file__).resolve().parent.parent.parent / "reference_docs"

_ROLE_PREAMBLE = """\
РОЛЬ И ГРАНИЦЫ
Ты — внутренний технический контролёр компании, выполняющей обследование зданий \
и сооружений различного назначения: промышленных, административных, жилых, \
объектов культурного наследия.
Твоя задача — выдать конкретный список замечаний, которые исполнитель может \
взять и исправить. Каждое замечание должно быть сформулировано так, чтобы было \
понятно: что не так, где находится, как исправить.
Ты никогда не:
- перепроверяешь инженерные расчёты
- оцениваешь правильность технических решений
- переписываешь текст целиком
- додумываешь содержание отсутствующих разделов — фиксируешь их отсутствие как замечание

Общий принцип: фиксируй только то, в чём уверен. Не уверен — не пиши."""

_OUTPUT_FORMAT = """\
ФОРМАТ ВЫДАЧИ
Выдавай пронумерованный список замечаний. Каждое замечание — по шаблону:

[Категория] № [номер]
Где: страница (из маркеров [Страница N] в тексте) / раздел / абзац
Что не так: конкретное описание ошибки
Как исправить: конкретное действие

ВАЖНО: Найди и выпиши ВСЕ замечания без исключения. Не сокращай список — \
проверяй документ от начала до конца. Полнота проверки важнее краткости.\
"""

REVIEW_PASSES = [
    {
        "needs_reference": False,
        "system_prompt": (
            f"{_ROLE_PREAMBLE}\n\n"
            "КАТЕГОРИИ ЗАМЕЧАНИЙ\n"
            "🔴 КРИТ — влияет на юридическую силу (адрес, заказчик, уровень ответственности, даты)\n"
            "🟠 ЛОГИКА — внутренние противоречия, несогласованность разделов\n"
            "⚪️ РЕКОМЕНДАЦИЯ — не ошибка, но стоит усилить (в области логики и критичных данных)\n\n"
            "ТИПИЧНЫЕ ОШИБКИ (искать в первую очередь)\n"
            "- Несовпадение адреса объекта в разных разделах\n"
            "- Несовпадение дат (договор / выезд / утверждение)\n"
            "- Противоречие в типе объекта: «здание» в одном месте, «фасады» в другом\n"
            "- Несогласованность цепочки: Ведомость дефектов → Заключение → Рекомендации\n"
            "- Категория технического состояния не вытекает из описанных дефектов\n\n"
            f"{_OUTPUT_FORMAT}"
        ),
    },
    {
        "needs_reference": False,
        "system_prompt": (
            f"{_ROLE_PREAMBLE}\n\n"
            "КАТЕГОРИИ ЗАМЕЧАНИЙ\n"
            "🟡 ШАБЛОН — остатки чужих данных, повторы\n"
            "🔵 ТЕКСТ — орфография, пунктуация, грамматика\n"
            "⚪️ РЕКОМЕНДАЦИЯ — не ошибка, но стоит усилить (в области текста и оформления)\n\n"
            "ТИПИЧНЫЕ ОШИБКИ (искать в первую очередь)\n"
            "- Остатки данных от предыдущего объекта (заказчик, название, габариты)\n"
            "- Разное написание одного термина по тексту\n\n"
            f"{_OUTPUT_FORMAT}"
        ),
    },
    {
        "needs_reference": True,
        "system_prompt": (
            f"{_ROLE_PREAMBLE}\n\n"
            "КАТЕГОРИИ ЗАМЕЧАНИЙ\n"
            "🟢 НОРМАТИВ — ошибки в ссылках на НТД\n"
            "⚪️ РЕКОМЕНДАЦИЯ — не ошибка, но стоит усилить (в области нормативных ссылок)\n\n"
            "НОРМАТИВНАЯ БАЗА\n"
            "Основные документы:\n"
            "- ГОСТ 31937-2024 — обследование и мониторинг зданий и сооружений\n"
            "- СП 13-102-2003 — обследование несущих строительных конструкций\n"
            "- ГОСТ 27751-2011 — надёжность строительных конструкций, уровни ответственности\n"
            "- ФЗ-384 — технический регламент, уровень ответственности\n"
            "- Актуальные СП по конструктивным разделам (СП 20, СП 22, СП 63 и др.)\n\n"
            "Ограничение: Проверяй только явные ошибки — например, ссылка на ГОСТ 31937-2011 "
            "вместо актуального ГОСТ 31937-2024, отменённый СНиП вместо действующего СП, "
            "неверный год редакции. В неочевидных случаях — пиши как рекомендацию. "
            "Не уверен — не пиши.\n\n"
            "ТИПИЧНЫЕ ОШИБКИ (искать в первую очередь)\n"
            "- Ссылка на ГОСТ 31937-2011 вместо ГОСТ 31937-2024\n"
            "- Ссылки на отменённые нормы без оговорок\n\n"
            f"{_OUTPUT_FORMAT}"
        ),
    },
]

MERGE_SYSTEM_PROMPT = """\
Ты объединяешь результаты нескольких проходов проверки технического отчёта в единый список.

Правила:
1. Удали дубликаты — одно и то же замечание, даже если сформулировано иначе.
2. Пронумеруй замечания последовательно.
3. Сохрани оригинальные формулировки замечаний.

ФОРМАТ ВЫДАЧИ
Каждое замечание — строго по шаблону:

[Категория] № [номер]
Где: страница / раздел / абзац
Что не так: конкретное описание ошибки
Как исправить: конкретное действие

Категории:
🔴 КРИТ — влияет на юридическую силу
🟠 ЛОГИКА — внутренние противоречия
🟡 ШАБЛОН — остатки чужих данных
🔵 ТЕКСТ — орфография, пунктуация, грамматика
🟢 НОРМАТИВ — ошибки в ссылках на НТД
⚪️ РЕКОМЕНДАЦИЯ — не ошибка, но стоит усилить

В конце — итоговая строка:
Всего замечаний: 🔴 _ / 🟠 _ / 🟡 _ / 🔵 _ / 🟢 _ / ⚪️ _
Если замечаний в категории нет — категорию не выводить.\
"""


class DocumentReviewService:
    """Сервис для парсинга и проверки технических отчётов."""

    def __init__(self):
        self._openai_client = None
        self._reference_texts: Optional[str] = None

    def _get_openai_client(self):
        if self._openai_client is None:
            from openai import OpenAI
            self._openai_client = OpenAI()
        return self._openai_client

    async def parse_document(self, document_name: str) -> str:
        """Скачивает документ из GCS и извлекает текст."""
        file_bytes, mime_type = await documents_storage.download(document_name)

        if document_name.lower().endswith(".pdf"):
            return self._parse_pdf(file_bytes)
        elif document_name.lower().endswith(".docx"):
            return self._parse_docx(file_bytes)
        else:
            raise ValueError(
                f"Неподдерживаемый формат файла: {document_name}. "
                f"Допустимые форматы: {', '.join(ALLOWED_EXTENSIONS)}"
            )

    async def review_document(
        self,
        document_name: str,
        prompt: Optional[str] = None,
        model: Optional[str] = None,
    ) -> dict:
        """Парсит документ и отправляет текст на проверку LLM."""
        model = model or DEFAULT_MODEL
        if model not in AVAILABLE_MODELS:
            raise ValueError(
                f"Неподдерживаемая модель: {model}. "
                f"Доступные: {', '.join(AVAILABLE_MODELS)}"
            )

        text = await self.parse_document(document_name)

        if not text.strip():
            return {
                "document_name": document_name,
                "extracted_text": "",
                "review_result": "Документ пуст или не содержит текста.",
            }

        review_result = await self._run_llm_review(text, prompt, model)

        return {
            "document_name": document_name,
            "extracted_text": text,
            "review_result": review_result,
        }

    # ── Парсинг ───────────────────────────────────────────────────

    @staticmethod
    def _has_page_break(paragraph) -> bool:
        """Проверяет наличие разрыва страницы в параграфе."""
        for run in paragraph.runs:
            br_elements = run._element.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br"
            )
            for br in br_elements:
                if br.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"
                ) == "page":
                    return True
        return False

    @staticmethod
    def _parse_docx(file_bytes: bytes) -> str:
        doc = DocxDocument(io.BytesIO(file_bytes))

        parts: list[str] = []
        page_num = 1
        parts.append(f"[Страница {page_num}]")

        for paragraph in doc.paragraphs:
            if DocumentReviewService._has_page_break(paragraph):
                page_num += 1
                parts.append(f"\n[Страница {page_num}]")
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)

    @staticmethod
    def _parse_pdf(file_bytes: bytes) -> str:
        doc = pymupdf.open(stream=file_bytes, filetype="pdf")
        parts: list[str] = []

        for page_num, page in enumerate(doc, start=1):
            text = page.get_text().strip()
            if text:
                parts.append(f"[Страница {page_num}]\n{text}")

        doc.close()
        return "\n".join(parts)

    # ── Reference docs (ГОСТы) ────────────────────────────────────

    def _load_reference_texts(self) -> str:
        """Загружает .txt файлы из reference_docs/ и склеивает."""
        if self._reference_texts is not None:
            return self._reference_texts

        if not REFERENCE_DOCS_DIR.exists():
            self._reference_texts = ""
            return self._reference_texts

        parts: list[str] = []
        for txt_file in sorted(REFERENCE_DOCS_DIR.glob("*.txt")):
            content = txt_file.read_text(encoding="utf-8").strip()
            if content:
                parts.append(f"=== {txt_file.stem} ===\n{content}")

        self._reference_texts = "\n\n".join(parts)
        if self._reference_texts:
            logger.info(
                "Загружено %d reference-документов (%d символов)",
                len(parts),
                len(self._reference_texts),
            )
        return self._reference_texts

    # ── LLM ───────────────────────────────────────────────────────

    async def _run_llm_review(
        self, text: str, prompt: Optional[str] = None, model: str = DEFAULT_MODEL,
    ) -> str:
        """Запускает параллельные фокусированные проходы и объединяет результаты."""
        reference = self._load_reference_texts()

        # Базовое сообщение без reference docs (проходы 1, 2)
        base_parts: list[str] = ["ТЕКСТ ОТЧЁТА ДЛЯ ПРОВЕРКИ:\n\n" + text]
        if prompt:
            base_parts.append("ДОПОЛНИТЕЛЬНЫЕ УКАЗАНИЯ:\n\n" + prompt)
        base_user_message = "\n\n---\n\n".join(base_parts)

        # Расширенное сообщение с reference docs (проход 3 — нормативы)
        if reference:
            ref_parts = [
                "НОРМАТИВНЫЕ ДОКУМЕНТЫ ДЛЯ СПРАВКИ:\n\n" + reference,
                "ТЕКСТ ОТЧЁТА ДЛЯ ПРОВЕРКИ:\n\n" + text,
            ]
            if prompt:
                ref_parts.append("ДОПОЛНИТЕЛЬНЫЕ УКАЗАНИЯ:\n\n" + prompt)
            ref_user_message = "\n\n---\n\n".join(ref_parts)
        else:
            ref_user_message = base_user_message

        logger.info(
            "Запуск %d параллельных проходов: %d символов текста, %d символов reference, модель %s",
            len(REVIEW_PASSES),
            len(text),
            len(reference),
            model,
        )

        tasks = [
            self._run_single_pass(
                ref_user_message if pass_cfg["needs_reference"] else base_user_message,
                pass_cfg,
                model,
            )
            for pass_cfg in REVIEW_PASSES
        ]
        pass_results = await asyncio.gather(*tasks)

        return await self._merge_results(pass_results, model)

    async def _call_llm(
        self, model: str, system_prompt: str, user_message: str,
    ) -> str:
        """Вызов OpenAI LLM."""
        client = self._get_openai_client()
        response = await asyncio.to_thread(
            client.chat.completions.create,
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
            temperature=REVIEW_TEMPERATURE,
            max_completion_tokens=REVIEW_MAX_TOKENS,
        )
        logger.info(
            "OpenAI [%s]: tokens %s/%s",
            model, response.usage.prompt_tokens, response.usage.completion_tokens,
        )
        return response.choices[0].message.content

    async def _run_single_pass(
        self, user_message: str, pass_cfg: dict, model: str,
    ) -> str:
        """Один фокусированный проход проверки."""
        try:
            result = await self._call_llm(model, pass_cfg["system_prompt"], user_message)
            logger.info("Проход завершён: %d символов", len(result))
            return result
        except Exception as e:
            logger.error("Ошибка прохода LLM: %s", e)
            raise

    async def _merge_results(self, pass_results: list[str], model: str) -> str:
        """Объединяет результаты проходов, удаляя дубликаты."""
        combined = "\n\n".join(
            f"--- Проход {i + 1} ---\n{result}"
            for i, result in enumerate(pass_results)
        )

        try:
            result = await self._call_llm(model, MERGE_SYSTEM_PROMPT, combined)
            logger.info("Объединение завершено: %d символов", len(result))
            return result
        except Exception as e:
            logger.error("Ошибка объединения результатов: %s", e)
            raise
