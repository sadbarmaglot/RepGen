"""
Генератор объединенных отчетов для Windows приложения
Объединяет ведомость дефектов, физический износ и конструктивные решения в один документ
"""

import sys
from pathlib import Path
from io import BytesIO
from datetime import datetime

# Добавляем корневую папку проекта в sys.path
root_dir = Path(__file__).parent.parent
sys.path.insert(0, str(root_dir))

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_unified_report(
    defects_data: list = None,
    wear_data: dict = None,
    constructive_data: list = None,
    object_name: str = "Объект обследования",
    address: str = "Адрес не указан",
    project_name: str = "Общий объект"
) -> BytesIO:
    """
    Создание объединенного отчета с ведомостью дефектов, физическим износом и конструктивными решениями
    
    Args:
        defects_data: Список результатов анализа дефектов
        wear_data: Данные расчета физического износа
        constructive_data: Список конструктивных решений
        object_name: Название объекта
        address: Адрес объекта
        project_name: Название объекта
        
    Returns:
        BytesIO: Буфер с документом
    """
    try:
        # Создаем новый документ
        doc = Document()
        
        # Заголовок документа
        title = doc.add_heading('ОТЧЕТ ПО ОБСЛЕДОВАНИЮ ТЕХНИЧЕСКОГО СОСТОЯНИЯ ЗДАНИЯ', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация об объекте
        info_para = doc.add_paragraph()
        info_para.add_run(f"Объект: ").bold = True
        info_para.add_run(object_name)
        info_para.add_run(f"\nАдрес: ").bold = True
        info_para.add_run(address)
        info_para.add_run(f"\nОбъект: ").bold = True
        info_para.add_run(project_name)
        info_para.add_run(f"\nДата обследования: ").bold = True
        info_para.add_run(datetime.now().strftime("%d.%m.%Y"))
        
        doc.add_paragraph()  # Пустая строка
        
        # 1. КОНСТРУКТИВНЫЕ РЕШЕНИЯ
        if constructive_data:
            doc.add_heading('1. КОНСТРУКТИВНЫЕ РЕШЕНИЯ', level=2)
            
            # Создаем таблицу конструктивных решений
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Заголовки таблицы
            header_cells = table.rows[0].cells
            headers = ['№ п/п', 'Конструкция', 'Описание']
            
            for i, header in enumerate(headers):
                header_cells[i].text = header
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(12)
                header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Заполняем таблицу данными
            for i, item in enumerate(constructive_data, 1):
                row_cells = table.add_row().cells
                
                row_cells[0].text = str(i)
                row_cells[1].text = item.get('category', '')
                row_cells[2].text = item.get('text', '')
                
                # Настраиваем размер шрифта
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(11)
            
            # Настраиваем ширину колонок
            table.columns[0].width = Cm(2)
            table.columns[1].width = Cm(4)
            table.columns[2].width = Cm(12)
            
            doc.add_paragraph()  # Пустая строка
        
        # 2. ФИЗИЧЕСКИЙ ИЗНОС
        if wear_data:
            doc.add_heading('2. РАСЧЕТ ФИЗИЧЕСКОГО ИЗНОСА ЗДАНИЯ', level=2)
            
            # Информация о методике
            method_para = doc.add_paragraph()
            method_para.add_run("Расчет выполнен в соответствии с ВСН 53-86р «Правила оценки физического износа жилых зданий»")
            method_para.italic = True
            
            doc.add_paragraph()  # Пустая строка
            
            # Создаем таблицу физического износа
            elements_data = wear_data.get('elements_table', [])
            if elements_data:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Заголовки таблицы
                header_cells = table.rows[0].cells
                headers = [
                    'Элементы здания', 
                    'Удельный вес, %',
                    'Физический износ, %', 
                    'Средневзвешенный износ'
                ]
                
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(11)
                    header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Заполняем таблицу данными
                for element in elements_data:
                    row_cells = table.add_row().cells
                    
                    row_cells[0].text = element.get('name', '')
                    row_cells[1].text = f"{element.get('weight', 0):.1f}"
                    row_cells[2].text = f"{element.get('wear', 0):.1f}"
                    row_cells[3].text = f"{element.get('weighted_wear', 0):.2f}"
                    
                    # Настраиваем размер шрифта и выравнивание
                    for j, cell in enumerate(row_cells):
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
                        if j > 0:  # Центрируем числовые данные
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Итоговая строка
                total_wear = wear_data.get('total_wear', 0)
                total_row = table.add_row().cells
                total_row[0].text = "ИТОГО"
                total_row[1].text = "100.0"
                total_row[2].text = "-"
                total_row[3].text = f"{total_wear:.1f}"
                
                # Делаем итоговую строку жирной
                for cell in total_row:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                    if cell != total_row[0]:  # Центрируем числовые данные
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Настраиваем ширину колонок
                table.columns[0].width = Cm(6)
                table.columns[1].width = Cm(3)
                table.columns[2].width = Cm(3)
                table.columns[3].width = Cm(4)
            
            # Заключение по износу
            doc.add_paragraph()
            conclusion_para = doc.add_paragraph()
            conclusion_para.add_run("Заключение: ").bold = True
            conclusion_para.add_run(f"Общий физический износ здания составляет {total_wear:.1f}%")
            
            # Определяем техническое состояние
            if total_wear <= 20:
                condition = "хорошее"
                recommendation = "Здание пригодно для эксплуатации"
            elif total_wear <= 40:
                condition = "удовлетворительное"
                recommendation = "Требуется текущий ремонт"
            elif total_wear <= 60:
                condition = "неудовлетворительное"
                recommendation = "Требуется капитальный ремонт"
            else:
                condition = "ветхое"
                recommendation = "Требуется реконструкция или снос"
            
            condition_para = doc.add_paragraph()
            condition_para.add_run("Техническое состояние: ").bold = True
            condition_para.add_run(condition)
            
            rec_para = doc.add_paragraph()
            rec_para.add_run("Рекомендации: ").bold = True
            rec_para.add_run(recommendation)
            
            doc.add_paragraph()  # Пустая строка
        
        # 3. ВЕДОМОСТЬ ДЕФЕКТОВ
        if defects_data:
            doc.add_heading('3. ВЕДОМОСТЬ ДЕФЕКТОВ И ПОВРЕЖДЕНИЙ', level=2)
            
            # Настраиваем ориентацию страницы для таблицы с фото
            section = doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
            
            # Создаем таблицу дефектов
            completed_results = [r for r in defects_data if r.get('analyzed', False)]
            
            if completed_results:
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                
                # Заголовки таблицы
                header_cells = table.rows[0].cells
                headers = ['№ п/п', 'Описание дефекта или повреждения', 
                          'Местоположение', 'Фотофиксация', 'Метод устранения']
                
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(12)
                
                # Заполняем таблицу данными
                for i, result in enumerate(completed_results, 1):
                    row_cells = table.add_row().cells
                    
                    row_cells[0].text = str(i)
                    row_cells[1].text = result.get('defect', 'Не определено')
                    row_cells[2].text = "По результатам анализа фотографии"
                    row_cells[4].text = result.get('eliminating_method', 'Не указан')
                    
                    # Вставляем фото в ячейку
                    photo_cell = row_cells[3]
                    photo_path = result.get('file_path', '')
                    
                    try:
                        if photo_path and Path(photo_path).exists():
                            # Очищаем ячейку
                            photo_cell.text = ""
                            
                            # Добавляем изображение
                            paragraph = photo_cell.paragraphs[0]
                            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                            
                            # Вставляем изображение с ограничением размера
                            run.add_picture(photo_path, width=Cm(4))
                            
                            # Центрируем изображение
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            photo_cell.text = f"Файл не найден: {result.get('filename', 'Неизвестно')}"
                            
                    except Exception as e:
                        print(f"Ошибка вставки изображения {photo_path}: {e}")
                        photo_cell.text = f"Ошибка загрузки: {result.get('filename', 'Неизвестно')}"
                    
                    # Настраиваем размер шрифта в текстовых ячейках
                    for j, cell in enumerate(row_cells):
                        if j != 3:  # Пропускаем ячейку с фото
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(11)
                    
                    # Настраиваем высоту строк для фотографий
                    row_cells[0]._parent.height = Cm(3.5)
                
                # Настраиваем ширину колонок
                widths = [Cm(1.5), Cm(5), Cm(3.5), Cm(8), Cm(5.5)]
                for i, width in enumerate(widths):
                    for row in table.rows:
                        row.cells[i].width = width
            else:
                doc.add_paragraph("Нет завершенных анализов дефектов для отображения.")
        
        # Возвращаем документ в буфере
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        print(f"Ошибка создания объединенного отчета: {e}")
        import traceback
        traceback.print_exc()
        return None

def save_unified_report(
    defects_data: list = None,
    wear_data: dict = None,
    constructive_data: list = None,
    output_path: str = None,
    object_name: str = "Объект обследования",
    address: str = "Адрес не указан",
    project_name: str = "Общий объект"
) -> bool:
    """
    Сохранение объединенного отчета в файл
    
    Args:
        defects_data: Список результатов анализа дефектов
        wear_data: Данные расчета физического износа
        constructive_data: Список конструктивных решений
        output_path: Путь для сохранения файла
        object_name: Название объекта
        address: Адрес объекта
        project_name: Название объекта
        
    Returns:
        bool: True если успешно
    """
    try:
        # Создаем документ
        buffer = create_unified_report(
            defects_data=defects_data,
            wear_data=wear_data,
            constructive_data=constructive_data,
            object_name=object_name,
            address=address,
            project_name=project_name
        )
        
        if buffer is None:
            return False
        
        # Сохраняем в файл
        with open(output_path, 'wb') as f:
            f.write(buffer.getvalue())
        
        return True
        
    except Exception as e:
        print(f"Ошибка сохранения объединенного отчета: {e}")
        return False

def generate_wear_report(report_data: dict, output_path: str) -> bool:
    """
    Создание отчета по физическому износу здания
    
    Args:
        report_data: Данные расчета износа
        output_path: Путь для сохранения файла
        
    Returns:
        bool: True если успешно
    """
    try:
        # Создаем новый документ
        doc = Document()
        
        # Настраиваем ориентацию страницы
        section = doc.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        
        # Заголовок документа
        title = doc.add_heading('РАСЧЕТ ФИЗИЧЕСКОГО ИЗНОСА ЗДАНИЯ', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Подзаголовок
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run(f"в соответствии с {report_data['vsn_reference']}")
        subtitle_run.italic = True
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Пустая строка
        
        # Информация о расчете
        info_para = doc.add_paragraph()
        info_para.add_run("Дата расчета: ").bold = True
        info_para.add_run(datetime.now().strftime("%d.%m.%Y"))
        
        doc.add_paragraph()  # Пустая строка
        
        # Создаем таблицу
        elements_data = report_data['elements_table']
        
        if elements_data:
            # Заголовок таблицы
            table_title = doc.add_paragraph()
            table_title_run = table_title.add_run("Таблица. Расчет физического износа конструктивных элементов здания")
            table_title_run.bold = True
            table_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Создаем таблицу
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Заголовки таблицы
            header_cells = table.rows[0].cells
            headers = [
                'Элементы здания', 
                'Удельный вес элемента в общей стоимости здания, %',
                'Физический износ элементов здания, %', 
                'Средневзвешенная степень физического износа'
            ]
            
            for i, header in enumerate(headers):
                header_cells[i].text = header
                # Делаем заголовки жирными
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
                header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Заполняем таблицу данными
            for element in elements_data:
                row_cells = table.add_row().cells
                
                row_cells[0].text = element['name']
                row_cells[1].text = f"{element['weight']:.1f}"
                row_cells[2].text = f"{element['wear']:.1f}"
                row_cells[3].text = f"{element['weighted_wear']:.2f}"
                
                # Настраиваем размер шрифта в ячейках
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                    # Центрируем числовые данные
                    if cell != row_cells[0]:  # Кроме первой колонки (названия)
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Итоговая строка
            total_row = table.add_row().cells
            total_row[0].text = "ИТОГО"
            total_row[1].text = f"{report_data['total_weight']:.1f}"
            total_row[2].text = "-"
            total_row[3].text = f"{report_data['total_wear']:.1f}"
            
            # Выделяем итоговую строку
            for cell in total_row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Настраиваем ширину колонок
            for i, width in enumerate([Cm(6), Cm(3.5), Cm(3.5), Cm(4)]):
                for row in table.rows:
                    row.cells[i].width = width
        
        doc.add_paragraph()  # Пустая строка
        
        # Заключение
        conclusion_title = doc.add_paragraph()
        conclusion_title_run = conclusion_title.add_run("ЗАКЛЮЧЕНИЕ")
        conclusion_title_run.bold = True
        conclusion_title_run.font.size = Pt(14)
        conclusion_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Результат расчета
        condition = report_data['condition']
        total_wear = report_data['total_wear']
        
        result_para = doc.add_paragraph()
        result_para.add_run("Физический износ здания в соответствии с ").font.size = Pt(12)
        result_para.add_run(f"{report_data['vsn_reference']} ").font.size = Pt(12)
        result_para.add_run("составляет ").font.size = Pt(12)
        
        wear_run = result_para.add_run(f"{total_wear}%")
        wear_run.bold = True
        wear_run.font.size = Pt(12)
        
        # Определяем цвет в зависимости от износа
        if total_wear <= 20:
            wear_run.font.color.rgb = RGBColor(0, 128, 0)  # Зеленый
        elif total_wear <= 40:
            wear_run.font.color.rgb = RGBColor(255, 140, 0)  # Оранжевый
        elif total_wear <= 60:
            wear_run.font.color.rgb = RGBColor(255, 0, 0)  # Красный
        else:
            wear_run.font.color.rgb = RGBColor(139, 0, 0)  # Темно-красный
        
        result_para.add_run(".").font.size = Pt(12)
        
        # Техническое состояние
        state_para = doc.add_paragraph()
        state_para.add_run("Техническое состояние здания: ").font.size = Pt(12)
        
        state_run = state_para.add_run(f"{condition['category'].lower()}")
        state_run.bold = True
        state_run.font.size = Pt(12)
        state_run.font.color.rgb = wear_run.font.color.rgb
        
        state_para.add_run(".").font.size = Pt(12)
        
        # Рекомендации
        recommendation_para = doc.add_paragraph()
        recommendation_para.add_run("Рекомендации: ").font.size = Pt(12)
        recommendation_para.add_run(f"{condition['recommendation'].lower()}.").font.size = Pt(12)
        
        doc.add_paragraph()  # Пустая строка
        
        # Методические указания
        method_para = doc.add_paragraph()
        method_para.add_run("Расчет выполнен в соответствии с методикой ").font.size = Pt(10)
        method_para.add_run(f"{report_data['vsn_reference']} ").font.size = Pt(10)
        method_para.add_run("\"Правила оценки физического износа жилых зданий\".").font.size = Pt(10)
        method_para.paragraph_format.space_before = Pt(20)
        
        # Подпись
        doc.add_paragraph()
        signature_para = doc.add_paragraph()
        signature_para.add_run("Расчет выполнен автоматически с использованием программного обеспечения.")
        signature_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        signature_para.paragraph_format.space_before = Pt(30)
        
        for run in signature_para.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Сохраняем документ
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Ошибка создания отчета по износу: {e}")
        return False
