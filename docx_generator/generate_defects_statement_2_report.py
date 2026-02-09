"""
Генератор отчета ведомости дефектов и повреждений №2.
Порт из десктоп-клиента для серверной генерации.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any, List, Dict, Optional

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT


def _set_cell_text(cell, text: str, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER, font_size: int = 10):
    """Устанавливает текст в ячейке с заданным форматированием."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text or "")
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.bold = bold


def _set_cell_image(cell, image_data: bytes, width=Cm(2.0)):
    """Вставляет изображение в ячейку таблицы."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(BytesIO(image_data), width=width)


def _normalize(value: Any) -> str:
    if value in (None, ""):
        return ""
    return str(value)


def generate_defects_statement_2_report(rows: List[Dict[str, Any]]) -> BytesIO:
    """
    Генерация DOCX-документа с ведомостью дефектов и повреждений №2.

    Args:
        rows: Список словарей с ключами:
            number, scheme_name, element_name, defect_volume,
            defects_and_causes, photo_data (bytes), danger_category,
            work_recommendations, recommended_work_types
    """
    doc = Document()

    # Альбомная ориентация
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Заголовок документа
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Ведомость дефектов и повреждений")
    title_run.bold = True
    title_run.font.size = Pt(14)

    doc.add_paragraph()

    # Создаем таблицу
    _create_defects_statement_2_table(doc, rows)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _create_defects_statement_2_table(doc: Document, rows_data: List[Dict[str, Any]]):
    row_count = len(rows_data)

    # Создаем таблицу: заголовок + строка нумерации + строки данных
    table = doc.add_table(rows=2 + row_count, cols=9)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    # Ширины столбцов (оптимизированы под альбомную ориентацию)
    widths = [
        Cm(1.0),   # № п/п
        Cm(2.5),   # Наименование схемы/изображения
        Cm(3.0),   # Наименование конструктивного элемента
        Cm(1.5),   # Объем дефектов
        Cm(4.0),   # Дефекты и вероятные причины
        Cm(2.5),   # Фотофиксация
        Cm(1.5),   # Категория опасности
        Cm(4.0),   # Рекомендации по составу работ
        Cm(4.0),   # Рекомендуемые виды работ
    ]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width

    # Заголовок (строка 0)
    header_row = table.rows[0]
    headers = [
        "№ п/п",
        "Наименование схемы/изображения",
        "Наименование конструктивного элемента/инженерной системы",
        "Объем дефектов",
        "Дефекты и вероятные причины возникновения",
        "Фотофиксация",
        "Категория опасности",
        "Рекомендации по составу работ",
        "Рекомендуемые виды работ"
    ]
    for idx, text in enumerate(headers):
        _set_cell_text(header_row.cells[idx], text, bold=True, font_size=10)

    # Строка с нумерацией столбцов (строка 1)
    numbering_row = table.rows[1]
    for idx in range(9):
        _set_cell_text(numbering_row.cells[idx], str(idx + 1), bold=True, font_size=10)

    # Данные (начиная со строки 2)
    for i, row_data in enumerate(rows_data):
        table_row = table.rows[2 + i]

        number = _normalize(row_data.get("number", ""))
        scheme_name = _normalize(row_data.get("scheme_name", ""))
        element_name = _normalize(row_data.get("element_name", ""))
        defect_volume = _normalize(row_data.get("defect_volume", ""))
        defects_and_causes = _normalize(row_data.get("defects_and_causes", ""))
        danger_category = _normalize(row_data.get("danger_category", ""))
        work_recommendations = _normalize(row_data.get("work_recommendations", ""))
        recommended_work_types = _normalize(row_data.get("recommended_work_types", ""))
        photo_data: Optional[bytes] = row_data.get("photo_data")

        _set_cell_text(table_row.cells[0], number, font_size=10)
        _set_cell_text(table_row.cells[1], scheme_name, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=10)
        _set_cell_text(table_row.cells[2], element_name, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=10)
        _set_cell_text(table_row.cells[3], defect_volume, font_size=10)
        _set_cell_text(table_row.cells[4], defects_and_causes, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=10)

        # Фото: вставляем изображение если есть данные, иначе пустая ячейка
        if photo_data:
            try:
                _set_cell_image(table_row.cells[5], photo_data, width=Cm(2.0))
            except Exception:
                _set_cell_text(table_row.cells[5], "", font_size=10)
        else:
            _set_cell_text(table_row.cells[5], "", font_size=10)

        _set_cell_text(table_row.cells[6], danger_category, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=10)
        _set_cell_text(table_row.cells[7], work_recommendations, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=10)
        _set_cell_text(table_row.cells[8], recommended_work_types, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=10)
