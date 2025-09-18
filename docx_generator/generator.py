from io import BytesIO
from docx import Document
from docx.enum.section import WD_SECTION, WD_ORIENT

from docx_generator.docx_utils import (
    add_header,
    add_footer,
    add_paragraph,
    add_run,
    constractive_decisions_table,
    defects_table,
    conclusion_table,
    wear_table,
    add_wear_text,
    add_wear_conclusion
)
from settings import (
    HEADER_TEXT,
    FOOTER_TEXT,
    COMMON_FONT,
    HEADER_FONT_SIZE,
    FOOTER_FONT_SIZE,
    ALIGN_CENTER,
    ALIGN_JUSTIFY,
    ALIGN_LEFT,
    TITLE_1,
    TITLE_1_1,
    BODY_TEXT_1_1,
    BODY_TEXT_1_2,
    BODY_TEXT_1_3,
    BODY_TEXT_1_4,
    TITLE_1_2,
    MAIN_FONT_SIZE,
    CD_TABLE_NUM_COLS,
    CD_TABLE_NUM_ROWS,
    DEFECTS_TABLE_NUM_COLS,
    DEFECTS_TABLE_NUM_ROWS,
    CONCLUSION_TABLE_NUM_COLS,
    CONCLUSION_TABLE_NUM_ROWS,
    TITLE_3,
    JOB_MAPPING,
    DEFAULT_WEAR_ELEMENTS,
    WEAR_TABLE_NUM_COLS,
    WEAR_TABLE_NUM_ROWS
)


def generate_full_report(user_data: dict) -> BytesIO:
    doc = Document()

    heading_style_1 = doc.styles['Body Text']
    heading_style_2 = doc.styles['Heading 2']
    body_style = doc.styles['Body Text']

    # Данные пользователя
    job_type = JOB_MAPPING[user_data["Page1"]["job_type"]]
    object_name = f'«{user_data["Page1"]["object_name"]}»'
    address_name = user_data["Page1"]["address_name"]
    footer_text = FOOTER_TEXT + address_name

    # Основной абзац
    body_text_1 = (
        BODY_TEXT_1_1 +
        job_type +
        BODY_TEXT_1_2 +
        object_name +
        BODY_TEXT_1_3 +
        address_name +
        BODY_TEXT_1_4
    )

    # Получаем данные конструктивных решений или используем данные по умолчанию
    user_constructions = user_data.get("Page3", {}).get("constructions", [])
    if user_constructions:
        constructions = [
            [row["construction"], row["description"]]
            for row in user_constructions
        ]
    else:
        constructions = user_data.get("constructive_decisions", DEFAULT_CONSTRUCTIVE_DECISIONS)

    defects = [
        [row.get("defect", ""), "", row.get("image_path", ""), row.get("eliminating_method", "")]
        for row in user_data.get("Page4", {}).get("defects", [])
    ]

    states = [
        ["", "", row["construction"], row["state"]]
        for row in user_data.get("Page5", {}).get("constructions", [])
    ]

    # Секция
    section = doc.sections[0]
    add_header(section, HEADER_TEXT, ALIGN_CENTER, HEADER_FONT_SIZE, COMMON_FONT)
    add_footer(section, footer_text, ALIGN_CENTER, FOOTER_FONT_SIZE, COMMON_FONT)

    # Заголовок 1
    head_1 = add_paragraph(doc=doc, style=heading_style_1, first_line_indent=0)
    add_run(head_1, TITLE_1, COMMON_FONT, MAIN_FONT_SIZE, bold=True)

    # Подзаголовок 1.1
    head_1_1 = add_paragraph(doc=doc, style=heading_style_2, first_line_indent=0)
    add_run(head_1_1, TITLE_1_1, COMMON_FONT, MAIN_FONT_SIZE, bold=True)

    # Основной текст
    body_1 = add_paragraph(doc=doc, style=body_style, alignment=ALIGN_JUSTIFY)
    add_run(body_1, body_text_1, COMMON_FONT, MAIN_FONT_SIZE, bold=False)

    # Подзаголовок 1.2
    head_1_2 = add_paragraph(doc=doc, style=heading_style_2, first_line_indent=0)
    add_run(head_1_2, TITLE_1_2, COMMON_FONT, MAIN_FONT_SIZE, bold=True)

    # Таблица конструктивных решений
    constractive_decisions_table(doc, constructions, CD_TABLE_NUM_COLS, CD_TABLE_NUM_ROWS)
    
    # Добавляем пустую строку
    doc.add_paragraph()
    
    # Добавляем раздел по физическому износу
    wear_title = add_paragraph(doc=doc, style=heading_style_2, first_line_indent=0)
    add_run(wear_title, "1.3. Оценка физического износа здания", COMMON_FONT, MAIN_FONT_SIZE, bold=True)
    
    # Добавляем текст о методике оценки
    add_wear_text(doc)
    
    # Таблица физического износа
    wear_data = user_data.get("wear_elements", DEFAULT_WEAR_ELEMENTS)
    wear_table(doc, wear_data, WEAR_TABLE_NUM_COLS, WEAR_TABLE_NUM_ROWS)
    
    # Добавляем заключение по износу
    total_wear = user_data.get("total_wear_percentage", 32)
    add_wear_conclusion(doc, total_wear)

    # Таблица дефектов (страница 4)
    if defects:
        new_section = doc.add_section(WD_SECTION.NEW_PAGE) # WD_SECTION.NEW_PAGE
        new_section.orientation = WD_ORIENT.LANDSCAPE

        # Меняем размеры страницы (ширина и высота меняются местами)
        new_width, new_height = new_section.page_height, new_section.page_width
        new_section.page_width = new_width
        new_section.page_height = new_height

        defects_table(doc, defects, DEFECTS_TABLE_NUM_COLS, DEFECTS_TABLE_NUM_ROWS)

    # Заключение / техническое состояние (страница 5)

    if states:
        new_section = doc.add_section(WD_SECTION.NEW_PAGE)  # WD_SECTION.NEW_PAGE
        new_section.orientation = WD_ORIENT.PORTRAIT

        head_3 = add_paragraph(doc=doc, style=heading_style_1, first_line_indent=0)
        add_run(head_3, TITLE_3, COMMON_FONT, MAIN_FONT_SIZE, bold=True)
        conclusion_table(doc, states, CONCLUSION_TABLE_NUM_COLS, CONCLUSION_TABLE_NUM_ROWS)

    # Сохранение в память
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
