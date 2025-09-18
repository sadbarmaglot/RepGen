import os
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Row, _Column
from docx.shared import Inches, Pt

from settings import (
    ALIGN_CENTER,
    ALIGN_LEFT,
    ALIGN_JUSTIFY,
    FIRST_LINE_INDENT,
    PAGE_LINE_SPACING,
    PAGE_RIGHT_INDENT,
    PAGE_LEFT_INDENT,
    PAGE_SPACE_BEFORE,
    TEXT_COLOR,
    CD_TABLE_HEAD_TEXTS,
    COMMON_FONT,
    CD_TABLE_HEAD_FONT_SIZE,
    CD_TABLE_COLUMN_WIDTHS,
    CD_TABLE_SUBHEAD_TEXTS,
    CD_TABLE_SUBHEAD_FONT_SIZE,
    DEFECTS_TABLE_HEAD_TEXTS,
    DEFECTS_TABLE_HEAD_FONT_SIZE,
    DEFECTS_TABLE_COLUMN_WIDTHS,
    DEFECTS_TABLE_SUBHEAD_TEXTS,
    CONCLUSION_TABLE_HEAD_TEXTS,
    CONCLUSION_TABLE_HEAD_FONT_SIZE,
    CONCLUSION_TABLE_COLUMN_WIDTHS,
    CONCLUSION_TABLE_SUBHEAD_TEXTS,
    WEAR_TABLE_HEAD_TEXTS,
    WEAR_TABLE_HEAD_FONT_SIZE,
    WEAR_TABLE_COLUMN_WIDTHS,
    WEAR_TABLE_SUBHEAD_TEXTS,
    WEAR_TABLE_SUBHEAD_FONT_SIZE
)


def insertHR(paragraph, position):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(
        pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement(f'w:{position}')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '10')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)


def add_header(section, header_text, align, font_size, font):
    header = section.header
    header_paragraph = header.paragraphs[0]

    header_paragraph.alignment = align
    header_paragraph_run = header_paragraph.add_run(header_text)
    header_paragraph_run.font.size = font_size
    header_paragraph_run.font.name = font
    insertHR(header_paragraph, "bottom")


def add_footer(section, footer_text, align, font_size, font):
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]

    footer_paragraph.alignment = align
    footer_paragraph_run = footer_paragraph.add_run(footer_text)
    footer_paragraph_run.font.size = font_size
    footer_paragraph_run.font.name = font
    insertHR(footer_paragraph, "top")


def add_paragraph(doc,
                  style,
                  alignment=ALIGN_CENTER,
                  first_line_indent=FIRST_LINE_INDENT,
                  line_spacing=PAGE_LINE_SPACING,
                  right_indent=PAGE_RIGHT_INDENT,
                  left_indent=PAGE_LEFT_INDENT,
                  space_before=PAGE_SPACE_BEFORE
                  ):
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = alignment
    paragraph.paragraph_format.first_line_indent = first_line_indent
    paragraph.paragraph_format.right_indent = right_indent
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.left_indent = left_indent
    paragraph.paragraph_format.space_before = space_before

    return paragraph


def add_run(paragraph, text, font, font_size, bold, font_color=TEXT_COLOR):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = font_size
    run.font.bold = bold
    run.font.color.rgb = font_color
    return run


def set_up_entity(entity, texts, font, font_size, bold, widths, align):
    if type(entity) in (_Row, _Column):
        for n, cell in enumerate(entity.cells):
            cell.width = widths[n]
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if "jpg" in texts[n] or "jpeg" in texts[n] or "png" in texts[n]:
                    run = paragraph.add_run()
                    run.add_picture(texts[n],  width=Inches(4))
                else:
                    paragraph.alignment = align
                    run = paragraph.add_run(texts[n])
                    run.font.name = font
                    run.font.size = font_size
                    run.font.bold = bold
    else:
        paragraphs = entity.paragraphs
        for paragraph in paragraphs:
            if "jpg" in texts[0] or "jpeg" in texts[0] or "png" in texts[0]:
                run = paragraph.add_run()
                run.add_picture(texts[0],  width=Inches(4))
            else:
                paragraph.alignment = align
                run = paragraph.add_run(texts[0])
                run.font.name = font
                run.font.size = font_size
                run.font.bold = bold


def constractive_decisions_table(doc, data, ncols, nrows, style="Table Grid"):
    # create new table
    table = doc.add_table(rows=nrows + len(data), cols=ncols, style=style)

    # set up head
    head_row = table.rows[0]
    # head_row.height = CD_TABLE_HEAD_HEIGHT
    set_up_entity(head_row,
                  CD_TABLE_HEAD_TEXTS,
                  COMMON_FONT,
                  CD_TABLE_HEAD_FONT_SIZE,
                  True,
                  CD_TABLE_COLUMN_WIDTHS,
                  ALIGN_CENTER
                  )

    # set up subhead
    subhead_row = table.rows[1]
    # subhead_row.height = CD_TABLE_SUBHEAD_HEIGHT

    # merge all cells in subhead
    subhead_row_cells = subhead_row.cells
    subhead_row_cell = subhead_row_cells[0].merge(subhead_row_cells[-1])
    set_up_entity(subhead_row_cell,
                  CD_TABLE_SUBHEAD_TEXTS,
                  COMMON_FONT,
                  CD_TABLE_SUBHEAD_FONT_SIZE,
                  True,
                  CD_TABLE_COLUMN_WIDTHS,
                  ALIGN_CENTER)

    # set up the rest rows
    for n, row in enumerate(table.rows[2:]):
        row_texts = [f"{n + 1}."] + data[n]
        set_up_entity(row,
                      row_texts,
                      COMMON_FONT,
                      CD_TABLE_SUBHEAD_FONT_SIZE,
                      False,
                      CD_TABLE_COLUMN_WIDTHS,
                      ALIGN_CENTER
                      )


def defects_table(doc, data, ncols, nrows, style="Table Grid"):
    # create new table
    table = doc.add_table(rows=nrows + len(data), cols=ncols, style=style)

    # set up head
    head_row = table.rows[0]
    # head_row.height = CD_TABLE_HEAD_HEIGHT
    set_up_entity(head_row,
                  DEFECTS_TABLE_HEAD_TEXTS,
                  COMMON_FONT,
                  DEFECTS_TABLE_HEAD_FONT_SIZE,
                  False,
                  DEFECTS_TABLE_COLUMN_WIDTHS,
                  ALIGN_CENTER
                  )

    # set up subhead
    subhead_row = table.rows[1]
    # subhead_row.height = CD_TABLE_SUBHEAD_HEIGHT
    set_up_entity(subhead_row,
                  DEFECTS_TABLE_SUBHEAD_TEXTS,
                  COMMON_FONT,
                  DEFECTS_TABLE_HEAD_FONT_SIZE,
                  False,
                  DEFECTS_TABLE_COLUMN_WIDTHS,
                  ALIGN_CENTER
                  )

    for i, row in enumerate(table.rows[2:]):
        defect_row = data[i]  # [description, "", image_path, recommendation]
        row_texts = [f"{i + 1}.", defect_row[0], defect_row[1], "", defect_row[3]]
        set_up_entity(row,
                      row_texts,
                      COMMON_FONT,
                      DEFECTS_TABLE_HEAD_FONT_SIZE,
                      False,
                      DEFECTS_TABLE_COLUMN_WIDTHS,
                      ALIGN_CENTER
                      )

        image_path = defect_row[2]
        if image_path and os.path.exists(image_path):
            try:
                run = row.cells[3].paragraphs[0].add_run()
                run.add_picture(image_path, width=Inches(4))
            except Exception as e:
                row.cells[3].text = "[Ошибка вставки фото]"
            finally:
                try:
                    os.remove(image_path)
                except Exception as e:
                    print(f"[WARNING] Не удалось удалить файл: {image_path} — {e}")
        else:
            row.cells[3].text = "📷 Файл не найден"


def conclusion_table(doc, data, ncols, nrows, style="Table Grid"):
    # create new table
    table = doc.add_table(rows=nrows + len(data), cols=ncols, style=style)

    # set up head
    head_row = table.rows[0]
    # head_row.height = CD_TABLE_HEAD_HEIGHT
    set_up_entity(head_row,
                  CONCLUSION_TABLE_HEAD_TEXTS,
                  COMMON_FONT,
                  CONCLUSION_TABLE_HEAD_FONT_SIZE,
                  False,
                  CONCLUSION_TABLE_COLUMN_WIDTHS,
                  ALIGN_LEFT
                  )
    for n, row in enumerate(table.rows[1:-1]):
        set_up_entity(row,
                      data[n],
                      COMMON_FONT,
                      CONCLUSION_TABLE_HEAD_FONT_SIZE,
                      False,
                      CONCLUSION_TABLE_COLUMN_WIDTHS,
                      ALIGN_LEFT
                      )

    fcolumn_cells = table.columns[0].cells
    fcolumn_cells[0].merge(fcolumn_cells[-2])

    fcolumn_cells = table.columns[1].cells
    fcolumn_cells[0].merge(fcolumn_cells[-2])

    # set up subhead
    subhead_row = table.rows[-1]
    # subhead_row.height = CD_TABLE_SUBHEAD_HEIGHT

    # merge all cells in subhead
    subhead_row_cells = subhead_row.cells
    subhead_row_cells[2].merge(subhead_row_cells[ncols - 1])

    set_up_entity(subhead_row,
                  CONCLUSION_TABLE_SUBHEAD_TEXTS,
                  COMMON_FONT,
                  CONCLUSION_TABLE_HEAD_FONT_SIZE,
                  False,
                  CONCLUSION_TABLE_COLUMN_WIDTHS,
                  ALIGN_LEFT
                  )


def wear_table(doc, data, ncols, nrows, style="Table Grid"):
    # Добавляем название таблицы
    table_title = doc.add_paragraph()
    table_title.alignment = ALIGN_CENTER
    table_title_run = table_title.add_run("Таблица П5.9.")
    table_title_run.font.name = COMMON_FONT
    table_title_run.font.size = Pt(14)
    table_title_run.font.bold = True
    
    # create new table
    table = doc.add_table(rows=nrows + len(data), cols=ncols, style=style)

    # set up head
    head_row = table.rows[0]
    set_up_entity(head_row,
                  WEAR_TABLE_HEAD_TEXTS,
                  COMMON_FONT,
                  WEAR_TABLE_HEAD_FONT_SIZE,
                  True,
                  WEAR_TABLE_COLUMN_WIDTHS,
                  ALIGN_CENTER
                  )

    # set up subhead
    subhead_row = table.rows[1]
    set_up_entity(subhead_row,
                  WEAR_TABLE_SUBHEAD_TEXTS,
                  COMMON_FONT,
                  WEAR_TABLE_SUBHEAD_FONT_SIZE,
                  False,
                  WEAR_TABLE_COLUMN_WIDTHS,
                  ALIGN_CENTER
                  )

    # merge cells for the third column header
    head_row.cells[2].merge(head_row.cells[3])

    # set up the rest rows
    for n, row in enumerate(table.rows[2:]):
        row_data = data[n]
        set_up_entity(row,
                      row_data,
                      COMMON_FONT,
                      WEAR_TABLE_SUBHEAD_FONT_SIZE,
                      False,
                      WEAR_TABLE_COLUMN_WIDTHS,
                      ALIGN_CENTER
                      )


def add_wear_text(doc):
    """Добавляет текст о методике оценки физического износа"""
    paragraph = doc.add_paragraph()
    paragraph.alignment = ALIGN_JUSTIFY
    paragraph.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    paragraph.paragraph_format.line_spacing = PAGE_LINE_SPACING
    
    text = ("Физический износ здания оценивается в соответствии с ВСН 53-86(р), который содержит таблицы показателей "
            "износа по каждому конструктивному элементу, виду инженерного оборудования и отделке. Эти показатели "
            "определяются при визуальном и инструментальном обследовании и проявляются в виде внешних признаков "
            "определенных дефектов на различных стадиях их развития.")
    
    run = paragraph.add_run(text)
    run.font.name = COMMON_FONT
    run.font.size = Pt(14)
    
    paragraph = doc.add_paragraph()
    paragraph.alignment = ALIGN_JUSTIFY
    paragraph.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    paragraph.paragraph_format.line_spacing = PAGE_LINE_SPACING
    
    text = ("Указанные проценты условно характеризуют степень физического износа как отношение стоимости ремонтных "
            "работ, необходимых для устранения дефектов, к стоимости замены элемента здания. Физический износ "
            "конструкции, элемента или системы с различной степенью износа на отдельных участках определяется как "
            "сумма показателей износа этих участков, взвешенная по их удельному весу в общем объеме соответствующих "
            "конструкций или элементов.")
    
    run = paragraph.add_run(text)
    run.font.name = COMMON_FONT
    run.font.size = Pt(14)
    
    paragraph = doc.add_paragraph()
    paragraph.alignment = ALIGN_JUSTIFY
    paragraph.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    paragraph.paragraph_format.line_spacing = PAGE_LINE_SPACING
    
    text = ("Общий физический износ здания (общий износ) определяется суммированием степеней износа его отдельных "
            "элементов, взвешенных по удельному весу их стоимости в общей восстановительной стоимости здания, "
            "что представлено в таблице П5.9.")
    
    run = paragraph.add_run(text)
    run.font.name = COMMON_FONT
    run.font.size = Pt(14)
    
    paragraph = doc.add_paragraph()
    paragraph.alignment = ALIGN_JUSTIFY
    paragraph.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    paragraph.paragraph_format.line_spacing = PAGE_LINE_SPACING
    
    text = ("Физический износ здания определяется с использованием удельных весов отдельных конструктивных элементов "
            "в процентах, установленных в приложении к таблице 3 раздела 1 (Лечебно-профилактические учреждения) "
            "Сборника №31 (Здания здравоохранения) «Укрупненные показатели восстановительной стоимости зданий для "
            "переоценки основных фондов». Капитальная группа I.")
    
    run = paragraph.add_run(text)
    run.font.name = COMMON_FONT
    run.font.size = Pt(14)


def add_wear_conclusion(doc, total_wear_percentage=32):
    """Добавляет заключение по физическому износу"""
    paragraph = doc.add_paragraph()
    paragraph.alignment = ALIGN_LEFT
    paragraph.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    paragraph.paragraph_format.line_spacing = PAGE_LINE_SPACING
    
    # Добавляем заголовок "Вывод"
    conclusion_title = paragraph.add_run("Вывод: ")
    conclusion_title.font.name = COMMON_FONT
    conclusion_title.font.size = Pt(14)
    conclusion_title.font.bold = True
    
    text = f"физический износ здания в соответствии с ВСН 53-86(р) составляет {total_wear_percentage}%. "
    if total_wear_percentage > 30:
        text += "Здание нуждается в проведении капитального ремонта."
    else:
        text += "Здание находится в удовлетворительном состоянии."
    
    run = paragraph.add_run(text)
    run.font.name = COMMON_FONT
    run.font.size = Pt(14)
    run.font.bold = False
