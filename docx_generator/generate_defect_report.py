from io import BytesIO
from docx import Document
from docx.enum.section import WD_ORIENT

from docx_generator.docx_utils import (
    defects_table, 
    constractive_decisions_table, 
    wear_table, 
    add_wear_text, 
    add_wear_conclusion
)
from settings import (
    DEFECTS_TABLE_NUM_COLS,
    DEFECTS_TABLE_NUM_ROWS,
    CD_TABLE_NUM_COLS,
    CD_TABLE_NUM_ROWS,
    WEAR_TABLE_NUM_COLS,
    WEAR_TABLE_NUM_ROWS,
    DEFAULT_CONSTRUCTIVE_DECISIONS,
    DEFAULT_WEAR_ELEMENTS
)

def generate_defect_only_report(user_data: dict) -> BytesIO:
    doc = Document()

    # Добавляем заголовок раздела конструктивных решений
    from docx_generator.docx_utils import add_paragraph, add_run
    from settings import TITLE_1_2, MAIN_FONT_SIZE, COMMON_FONT, TEXT_COLOR
    
    # Заголовок раздела конструктивных решений
    title_paragraph = add_paragraph(doc, "Normal")
    add_run(title_paragraph, TITLE_1_2, COMMON_FONT, MAIN_FONT_SIZE, True, TEXT_COLOR)
    
    # Таблица конструктивных решений
    constructive_data = user_data.get("constructive_decisions", DEFAULT_CONSTRUCTIVE_DECISIONS)
    constractive_decisions_table(doc, constructive_data, CD_TABLE_NUM_COLS, CD_TABLE_NUM_ROWS)
    
    # Добавляем пустую строку
    doc.add_paragraph()
    
    # Добавляем заголовок раздела физического износа
    wear_title_paragraph = add_paragraph(doc, "Normal")
    add_run(wear_title_paragraph, "1.3. Оценка физического износа здания", COMMON_FONT, MAIN_FONT_SIZE, True, TEXT_COLOR)
    
    # Добавляем текст о методике оценки
    add_wear_text(doc)
    
    # Таблица физического износа
    wear_data = user_data.get("wear_elements", DEFAULT_WEAR_ELEMENTS)
    wear_table(doc, wear_data, WEAR_TABLE_NUM_COLS, WEAR_TABLE_NUM_ROWS)
    
    # Добавляем заключение по износу
    total_wear = user_data.get("total_wear_percentage", 32)
    add_wear_conclusion(doc, total_wear)
    
    # Добавляем пустую строку
    doc.add_paragraph()
    
    # Добавляем заголовок раздела дефектов
    defects_title_paragraph = add_paragraph(doc, "Normal")
    add_run(defects_title_paragraph, "2. Ведомость дефектов и повреждений", COMMON_FONT, MAIN_FONT_SIZE, True, TEXT_COLOR)

    defects = [
        [
            row.get("defect", ""),
            "",
            row.get("image_path", ""),
            row.get("eliminating_method", "")
        ]
        for row in user_data.get("defects", [])
    ]

    if defects:
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

        defects_table(doc, defects, DEFECTS_TABLE_NUM_COLS, DEFECTS_TABLE_NUM_ROWS)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
